"""DOCX text extraction via python-docx."""

from __future__ import annotations

from io import BytesIO

from docx import Document

from app.extractors.base import ExtractedText


def extract_docx_text(data: bytes) -> ExtractedText:
    """Flatten paragraphs + tables into a single plain-text blob."""
    document = Document(BytesIO(data))

    parts: list[str] = [p.text for p in document.paragraphs if p.text and p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return ExtractedText(
        text="\n".join(parts).strip(),
        source="docx_text",
        page_count=1,
    )
